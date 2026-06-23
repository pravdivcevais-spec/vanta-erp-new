"""
docs.py — генерация документов для логистических операций.

Функции возвращают bytes (docx), которые можно передать
напрямую в st.download_button(data=...).
"""

from __future__ import annotations

import io
from datetime import datetime


def generate_logistics_doc_bytes(
    logistics_row: dict,
    bike_details: list[dict],
    darkstore_name: str = "",
    created_by_name: str = "",
    master_name: str = "",
    vyvoz_bikes: list[dict] | None = None,
    postavka_bikes: list[dict] | None = None,
) -> bytes:
    """
    Генерирует акт приёма-передачи велосипедов в формате .docx.

    Параметры:
        logistics_row   — строка из logistics_request
        bike_details    — список dict с ключами: gov_number, serial_number, model, tech_status
                          (для вывоза/поставки; для замены — не используется)
        darkstore_name  — название даркстора
        created_by_name — ФИО диспетчера
        master_name     — ФИО выездного мастера
        vyvoz_bikes     — байки на вывоз (только для типа "замена")
        postavka_bikes  — байки на поставку (только для типа "замена")

    Возвращает bytes для st.download_button.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError:
        raise RuntimeError(
            "Для генерации документов установите: pip install python-docx"
        )

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)

    rtype = logistics_row.get("request_type", "")
    if rtype == "поставка":
        rtype_label = "ПОСТАВКИ"
    elif rtype == "замена":
        rtype_label = "ЗАМЕНЫ"
    else:
        rtype_label = "ВЫВОЗА"
    doc_date = datetime.now().strftime("%d.%m.%Y")
    doc_number = logistics_row.get("id", "—")

    # ---------------------------------------------------------------------------
    # Заголовок
    # ---------------------------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"АКТ {rtype_label} ВЕЛОСИПЕДОВ")
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"№ {doc_number}  от  {doc_date}").font.size = Pt(11)

    doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # Реквизиты
    # ---------------------------------------------------------------------------
    def _field(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value or "—")
        p.paragraph_format.space_after = Pt(2)

    _field("Даркстор", darkstore_name or logistics_row.get("darkstore_name") or "—")
    op_label = {"поставка": "Поставка велосипедов", "замена": "Замена велосипедов"}.get(rtype, "Вывоз велосипедов")
    _field("Тип операции", op_label)
    _field("Ответственный диспетчер", created_by_name or "—")
    _field("Выездной мастер", master_name or (
        (logistics_row.get("master_first_name", "") + " " + logistics_row.get("master_last_name", "")).strip()
    ) or "—")
    if logistics_row.get("notes"):
        _field("Примечание", logistics_row["notes"])

    doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # Таблица велосипедов
    # ---------------------------------------------------------------------------
    def _bike_table(header_label: str, bikes: list[dict]) -> None:
        p_hdr = doc.add_paragraph()
        p_hdr.add_run(header_label).bold = True

        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"

        col_widths = [Cm(1.2), Cm(3.0), Cm(6.0), Cm(4.5), Cm(4.0)]
        for i, width in enumerate(col_widths):
            for cell in tbl.columns[i].cells:
                cell.width = width

        col_headers = ["№", "Гос номер", "Модель", "Серийный номер", "Состояние"]
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(col_headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        for idx, bike in enumerate(bikes, start=1):
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = str(bike.get("gov_number") or "—")
            row_cells[2].text = str(bike.get("model") or "—")
            row_cells[3].text = str(bike.get("serial_number") or "—")
            row_cells[4].text = str(bike.get("tech_status") or "—")

        doc.add_paragraph()

    if rtype == "замена":
        _bike_table("Велосипеды на вывоз (старые, с даркстора):", vyvoz_bikes or [])
        _bike_table("Велосипеды на поставку (новые, со склада):", postavka_bikes or [])
        total_count = len(vyvoz_bikes or []) + len(postavka_bikes or [])
    else:
        _bike_table("Перечень велосипедов:", bike_details)
        total_count = len(bike_details)

    # ---------------------------------------------------------------------------
    # Итог
    # ---------------------------------------------------------------------------
    total_p = doc.add_paragraph()
    total_p.add_run("Итого велосипедов: ").bold = True
    total_p.add_run(str(total_count))

    doc.add_paragraph()
    doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # Подписи
    # ---------------------------------------------------------------------------
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = "Table Grid"

    def _sig_cell(cell, label: str, name: str = "") -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.add_run(f"{label}:\n").bold = True
        p.add_run(f"{name or '________________________'}\n\n")
        p.add_run("Подпись: ________________________")

    _sig_cell(sig_table.rows[0].cells[0], "Диспетчер",     created_by_name)
    _sig_cell(sig_table.rows[0].cells[1], "Мастер",        master_name)
    _sig_cell(sig_table.rows[1].cells[0], "Сдал (даркстор)", "")
    if rtype == "вывоз":
        recv_label = "Принял (склад)"
    elif rtype == "замена":
        recv_label = "Принял (даркстор / склад)"
    else:
        recv_label = "Принял (даркстор)"
    _sig_cell(sig_table.rows[1].cells[1], recv_label, "")
    sig_table.rows[2].cells[0].text = f"Дата: {doc_date}"
    sig_table.rows[2].cells[1].text = f"Дата: {doc_date}"

    # ---------------------------------------------------------------------------
    # Сохранение в bytes
    # ---------------------------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def logistics_doc_filename(logistics_row: dict) -> str:
    """Возвращает имя файла для скачивания."""
    rtype_map = {"поставка": "postavka", "замена": "zamena"}
    rtype = rtype_map.get(logistics_row.get("request_type", ""), "vyvoz")
    lid = logistics_row.get("id", "0")
    date_str = datetime.now().strftime("%Y%m%d")
    return f"akt_{rtype}_{lid}_{date_str}.docx"
